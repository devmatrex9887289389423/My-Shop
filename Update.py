import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import csv
from datetime import datetime
import time
import re

try:
    import docx
except ImportError:
    pass

class CraftsmanProSystem:
    def __init__(self, root):
        self.root = root
        self.root.title("Craftsman System PRO")
        self.root.geometry("500x400")
        
        # ملفات السيستم الثابتة
        self.goods_file = "السلع_والمخزن.csv"
        self.sales_file = "عمليات_البيع.csv"
        self.users_file = "قاعدة_المشرفين.csv"
        self.config_file = "اعدادات_النظام.csv"
        
        self.cart = [] 
        self.current_user = "" 
        self.users_db = {}
        
        self.current_lang = "ar"
        self.current_theme = "Dark"
        
        self.themes = {
            "Dark": {"bg": "#1E2229", "card": "#282C34", "text": "white", "accent": "#00ADB5", "btn_save": "#2ECC71"},
            "Light": {"bg": "#F5F5F5", "card": "#FFFFFF", "text": "#333333", "accent": "#2980B9", "btn_save": "#27AE60"},
            "Blue": {"bg": "#0B132B", "card": "#1C2541", "text": "#FFFFFF", "accent": "#5BC0BE", "btn_save": "#3A7CA5"},
            "Green": {"bg": "#1A2E1A", "card": "#243E24", "text": "#E8F5E9", "accent": "#A9DFBF", "btn_save": "#2E7D32"}
        }
        
        self.lang_pack = {
            "ar": {
                "login_title": "🔒 نظام تسجيل دخول المشرفين", "choose_user": "اختر اسم المشرف:",
                "enter_pass": "ادخل كلمة المرور الخاصة بك:", "login_btn": "دخول السيستم 🚀",
                "supervisor": "👤 المشرف الحالي: ", "settings_btn": "⚙️ الإعدادات والمشرفين",
                "search_title": " 🔍 البحث عن السلع بالشبه في المخزن ", "table_title": " 📦 جرد البضاعة المتاحة حالياً ",
                "code": "كود المقاس", "name": "اسم السلعة", "price": "السعر", "stock": "المخزن",
                "add_manual": " 📥 إضافة صنف جديد للمخزن يدوياً ", "save_btn": "حفظ وتثبيت الصنف في إكسل السلع 💾",
                "cart_title": " 🧾 سلة مبيعات الفاتورة الحاضرة ", "sell_place": "📍 مكان البيع الحالي:",
                "size": "المقاس:", "qty": "الكمية:", "add_cart": "أضف للفاتورة ➕", "total": "الإجمالي",
                "invoice_total": "إجمالي الفاتورة: ", "checkout_btn": "إتمام وترحيل الفاتورة بالكامل 💾",
                "history_title": " 🔍 سيرش وتصفية الفواتير ", "time": "الوقت", "account": "الحساب",
                "settings_win": "⚙️ لوحة الإعدادات والمشرفين واللغات", "manage_users": " 👥 إدارة حسابات وباسووردات المشرفين ",
                "user_name": "اسم المشرف:", "new_pass": "الباسوورد الجديدة:", "save_user_btn": "💾 إضافة أو تغيير باسوورد المشرف",
                "del_user_btn": "❌ حذف مشرف من السيستم", "word_title": " 📝 استيراد بضاعة من ملف Word (منقذ البيانات) ",
                "word_btn": "📂 اختيار ملف الوورد وقراءته فوراً (.docx)", "lang_label": "🌐 لغة السيستم (Language):",
                "theme_label": "🎨 ثيم السيستم (Theme):", "currency": " ج.م"
            },
            "en": {
                "login_title": "🔒 Supervisor Login System", "choose_user": "Select Supervisor:",
                "enter_pass": "Enter Your Password:", "login_btn": "Login to System 🚀",
                "supervisor": "👤 Current User: ", "settings_btn": "⚙️ Settings & Users",
                "search_title": " 🔍 Search Products in Inventory ", "table_title": " 📦 Current Available Stock ",
                "code": "Size Code", "name": "Item Name", "price": "Price", "stock": "Stock",
                "add_manual": " 📥 Add New Item Manually ", "save_btn": "Save Item to Excel 💾",
                "cart_title": " 🧾 Current Sales Invoice Cart ", "sell_place": "📍 Sales Location:",
                "size": "Size:", "qty": "Qty:", "add_cart": "Add to Cart ➕", "total": "Total",
                "invoice_total": "Invoice Total: ", "checkout_btn": "Checkout & Archive Invoice 💾",
                "history_title": " 🔍 Search & Filter Invoices ", "time": "Time", "account": "Amount",
                "settings_win": "⚙️ Settings, Users & Languages", "manage_users": " 👥 Manage Supervisor Accounts ",
                "user_name": "Username:", "new_pass": "New Password:", "save_user_btn": "💾 Add/Update Supervisor",
                "del_user_btn": "❌ Delete Supervisor", "word_title": " 📝 Import Inventory from Word File ",
                "word_btn": "📂 Choose Word File (.docx)", "lang_label": "Language:",
                "theme_label": "🎨 System Theme:", "currency": " EGP"
            }
        }
        
        self.init_files()
        self.load_config()
        self.load_users()
        self.apply_theme_to_root()
        self.show_login_screen()

    def init_files(self):
        if not os.path.exists(self.goods_file):
            with open(self.goods_file, "w", newline="", encoding="utf-8-sig") as f:
                csv.writer(f).writerow(["الكود", "اسم السلعة", "السعر", "المخزن"])
        if not os.path.exists(self.sales_file):
            with open(self.sales_file, "w", newline="", encoding="utf-8-sig") as f:
                csv.writer(f).writerow(["التاريخ والوقت", "مكان البيع", "المشرف المسؤول", "الكود", "الاسم", "الكمية المباعة", "الإجمالي"])
        if not os.path.exists(self.users_file):
            with open(self.users_file, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["المشرف", "الرقم السري"])
                writer.writerow(["Ryder", "2026"])
                writer.writerow(["MatrixYasta", "2026"])
        if not os.path.exists(self.config_file):
            with open(self.config_file, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["اللغة", "الثيم"])
                writer.writerow(["ar", "Dark"])

    def load_config(self):
        if os.path.exists(self.config_file):
            with open(self.config_file, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                try:
                    next(reader)
                    row = next(reader)
                    if row and len(row) >= 2:
                        self.current_lang = row[0]
                        self.current_theme = row[1]
                except: pass

    def save_config(self):
        with open(self.config_file, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["اللغة", "الثيم"])
            writer.writerow([self.current_lang, self.current_theme])

    def load_users(self):
        self.users_db = {}
        if os.path.exists(self.users_file):
            with open(self.users_file, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                try:
                    next(reader)
                    for row in reader:
                        if row and len(row) >= 2: self.users_db[row[0]] = row[1]
                except StopIteration: pass

    def save_users(self):
        with open(self.users_file, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["المشرف", "الرقم السري"])
            for user, password in self.users_db.items(): writer.writerow([user, password])

    def load_goods(self):
        goods = {}
        if os.path.exists(self.goods_file):
            with open(self.goods_file, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                try:
                    next(reader)
                    for row in reader:
                        if row and len(row) >= 4: goods[row[0]] = {"name": row[1], "price": float(row[2]), "stock": int(row[3])}
                except StopIteration: pass
        return goods

    def save_goods(self, goods):
        with open(self.goods_file, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerow(["الكود", "اسم السلعة", "السعر", "المخزن"])
            for code, info in goods.items(): writer.writerow([code, info["name"], info["price"], info["stock"]])

    def apply_theme_to_root(self):
        t = self.themes[self.current_theme]
        self.root.configure(bg=t["bg"])

    def show_login_screen(self):
        t = self.themes[self.current_theme]
        l = self.lang_pack[self.current_lang]
        
        self.login_frame = tk.Frame(self.root, bg=t["bg"])
        self.login_frame.pack(fill="both", expand=True, pady=30)
        
        tk.Label(self.login_frame, text=l["login_title"], font=("Segoe UI", 14, "bold"), bg=t["bg"], fg=t["accent"]).pack(pady=10)
        tk.Label(self.login_frame, text=l["choose_user"], font=("Segoe UI", 10), bg=t["bg"], fg=t["text"]).pack(pady=2)
        
        user_list = list(self.users_db.keys()) if self.users_db else ["Ryder", "MatrixYasta"]
        self.combo_user = ttk.Combobox(self.login_frame, values=user_list, font=("Segoe UI", 11), state="readonly", width=18)
        self.combo_user.set(user_list[0])
        self.combo_user.pack(pady=5)
        
        tk.Label(self.login_frame, text=l["enter_pass"], font=("Segoe UI", 10), bg=t["bg"], fg=t["text"]).pack(pady=2)
        self.pass_entry = tk.Entry(self.login_frame, font=("Segoe UI", 11), show="*", justify="center", width=20)
        self.pass_entry.pack(pady=5)
        
        btn_login = tk.Button(self.login_frame, text=l["login_btn"], font=("Segoe UI", 11, "bold"), bg="#2ECC71", fg="white", width=15, bd=0, command=self.check_password)
        btn_login.pack(pady=20)

    def check_password(self):
        selected_user = self.combo_user.get()
        entered_pass = self.pass_entry.get()
        if selected_user in self.users_db and self.users_db[selected_user] == entered_pass:
            self.current_user = selected_user
            self.login_frame.destroy()
            self.show_main_dashboard()
        else:
            messagebox.showerror("Error", "Wrong Password!")
            self.pass_entry.delete(0, "end")

    def show_main_dashboard(self):
        self.root.geometry("1280x780")
        self.inventory = self.load_goods()
        self.refresh_dashboard_ui()

    def refresh_dashboard_ui(self):
        for widget in self.root.winfo_children():
            widget.destroy()
            
        t = self.themes[self.current_theme]
        l = self.lang_pack[self.current_lang]
        self.apply_theme_to_root()
        
        top_bar = tk.Frame(self.root, bg=t["card"], height=45)
        top_bar.pack(fill="x", side="top")
        top_bar.pack_propagate(False)
        
        side_align = "right" if self.current_lang == "ar" else "left"
        opposite_align = "left" if self.current_lang == "ar" else "right"
        
        lbl_user_status = tk.Label(top_bar, text=f"{l['supervisor']}{self.current_user}", font=("Segoe UI", 11, "bold"), bg=t["card"], fg="#2ECC71")
        lbl_user_status.pack(side=side_align, padx=15, pady=10)
        
        self.lbl_clock = tk.Label(top_bar, text="", font=("Segoe UI", 12, "bold"), bg=t["card"], fg="#FFD700")
        self.lbl_clock.pack(side=side_align, padx=30, pady=10)
        self.update_live_clock()
        
        btn_settings = tk.Button(top_bar, text=l["settings_btn"], font=("Segoe UI", 10, "bold"), bg="#E74C3C", fg="white", bd=0, padx=10, command=self.open_settings_window)
        btn_settings.pack(side=opposite_align, padx=15, pady=8)
        
        window_container = tk.Frame(self.root, bg=t["bg"])
        window_container.pack(fill="both", expand=True, padx=15, pady=10)
        
        right_side = tk.Frame(window_container, bg=t["bg"])
        right_side.pack(side=side_align, fill="both", expand=True, padx=10)
        
        left_side = tk.Frame(window_container, bg=t["bg"])
        left_side.pack(side=opposite_align, fill="both", expand=True, padx=10)
        
        goods_search_card = tk.LabelFrame(right_side, text=l["search_title"], font=("Segoe UI", 11, "bold"), bg=t["card"], fg=t["accent"], bd=1)
        goods_search_card.pack(fill="x", pady=5, ipady=5)
        self.search_entry = tk.Entry(goods_search_card, font=("Segoe UI", 11), bg=t["bg"], fg=t["text"])
        self.search_entry.pack(fill="x", padx=10, pady=5)
        self.search_entry.bind("<KeyRelease>", self.filter_inventory)
        
        inventory_table_card = tk.LabelFrame(right_side, text=l["table_title"], font=("Segoe UI", 11, "bold"), bg=t["card"], fg=t["accent"], bd=1)
        inventory_table_card.pack(fill="both", expand=True, pady=5)
        
        self.tree_goods = ttk.Treeview(inventory_table_card, columns=("code", "name", "price", "stock"), show="headings")
        self.tree_goods.heading("code", text=l["code"])
        self.tree_goods.heading("name", text=l["name"])
        self.tree_goods.heading("price", text=l["price"])
        self.tree_goods.heading("stock", text=l["stock"])
        self.tree_goods.column("code", anchor="center", width=80)
        self.tree_goods.column("name", anchor="center" if self.current_lang=="en" else "e", width=200)
        self.tree_goods.column("price", anchor="center", width=80)
        self.tree_goods.column("stock", anchor="center", width=80)
        self.tree_goods.pack(fill="both", expand=True, padx=10, pady=10)
        
        add_product_card = tk.LabelFrame(right_side, text=l["add_manual"], font=("Segoe UI", 11, "bold"), bg=t["card"], fg="#2ECC71", bd=1)
        add_product_card.pack(fill="x", pady=5, ipady=5)
        
        f_row = tk.Frame(add_product_card, bg=t["card"])
        f_row.pack(fill="x", padx=10, pady=5)
        
        tk.Label(f_row, text=l["code"]+":", font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.new_code = tk.Entry(f_row, width=8)
        self.new_code.pack(side=side_align, padx=3)
        
        tk.Label(f_row, text=l["name"]+":", font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.new_name = tk.Entry(f_row, width=14)
        self.new_name.pack(side=side_align, padx=3)
        
        tk.Label(f_row, text=l["price"]+":", font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.new_price = tk.Entry(f_row, width=6)
        self.new_price.pack(side=side_align, padx=3)
        
        tk.Label(f_row, text=l["stock"]+":", font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.new_stock = tk.Entry(f_row, width=6)
        self.new_stock.pack(side=side_align, padx=3)
        
        btn_save = tk.Button(add_product_card, text=l["save_btn"], bg=t["btn_save"], fg="white", font=("Segoe UI", 10, "bold"), bd=0, command=self.add_custom_product)
        btn_save.pack(fill="x", padx=10, pady=5)

        invoice_card = tk.LabelFrame(left_side, text=l["cart_title"], font=("Segoe UI", 11, "bold"), bg=t["card"], fg="#E67E22", bd=1)
        invoice_card.pack(fill="both", expand=True, pady=5)
        
        loc_panel = tk.Frame(invoice_card, bg=t["card"])
        loc_panel.pack(fill="x", padx=10, pady=5)
        tk.Label(loc_panel, text=l["sell_place"], font=("Segoe UI", 10, "bold"), bg=t["card"], fg=t["accent"]).pack(side=side_align, padx=5)
        self.combo_location = ttk.Combobox(loc_panel, values=["المحل الرئيسي", "المخزن العلوي", "الفرع الثاني"], width=15, state="readonly")
        self.combo_location.set("المحل الرئيسي")
        self.combo_location.pack(side=side_align, padx=5)
        
        item_panel = tk.Frame(invoice_card, bg=t["card"])
        item_panel.pack(fill="x", padx=10, pady=5)
        tk.Label(item_panel, text=l["size"], font=("Segoe UI", 10), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.combo_size = ttk.Combobox(item_panel, values=list(self.inventory.keys()), width=10, state="readonly")
        self.combo_size.pack(side=side_align, padx=2)
        
        tk.Label(item_panel, text=l["qty"], font=("Segoe UI", 10), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        self.qty_entry = tk.Entry(item_panel, width=6)
        self.qty_entry.pack(side=side_align, padx=2)
        
        btn_add_cart = tk.Button(item_panel, text=l["add_cart"], bg="#E67E22", fg="white", font=("Segoe UI", 10, "bold"), bd=0, command=self.add_to_cart)
        btn_add_cart.pack(side=side_align, padx=10)
        
        self.tree_cart = ttk.Treeview(invoice_card, columns=("code", "name", "qty", "total"), show="headings", height=5)
        self.tree_cart.heading("code", text=l["code"])
        self.tree_cart.heading("name", text=l["name"])
        self.tree_cart.heading("qty", text=l["qty"])
        self.tree_cart.heading("total", text=l["total"])
        self.tree_cart.column("code", anchor="center", width=60)
        self.tree_cart.column("name", anchor="center" if self.current_lang=="en" else "e", width=140)
        self.tree_cart.column("qty", anchor="center", width=50)
        self.tree_cart.column("total", anchor="center", width=70)
        self.tree_cart.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.lbl_cart_total = tk.Label(invoice_card, text=l["invoice_total"]+f"0.00{l['currency']}", font=("Segoe UI", 12, "bold"), bg=t["card"], fg="#FFD700")
        self.lbl_cart_total.pack(pady=2)
        
        btn_checkout = tk.Button(invoice_card, text=l["checkout_btn"], bg="#D35400", fg="white", font=("Segoe UI", 11, "bold"), bd=0, height=2, command=self.checkout_cart)
        btn_checkout.pack(fill="x", padx=10, pady=5)

        history_search_card = tk.LabelFrame(left_side, text=l["history_title"], font=("Segoe UI", 11, "bold"), bg=t["card"], fg="#FFD700", bd=1)
        history_search_card.pack(fill="x", pady=5, ipady=5)
        
        self.tree_sales_history = ttk.Treeview(history_search_card, columns=("time", "loc", "user", "code", "qty", "total"), show="headings", height=5)
        self.tree_sales_history.heading("time", text=l["time"])
        self.tree_sales_history.heading("loc", text="Location" if self.current_lang=="en" else "مكان البيع")
        self.tree_sales_history.heading("user", text="User" if self.current_lang=="en" else "المشرف")
        self.tree_sales_history.heading("code", text=l["code"])
        self.tree_sales_history.heading("qty", text=l["qty"])
        self.tree_sales_history.heading("total", text=l["account"])
        self.tree_sales_history.column("time", anchor="center", width=110)
        self.tree_sales_history.column("loc", anchor="center", width=80)
        self.tree_sales_history.column("user", anchor="center", width=80)
        self.tree_sales_history.column("code", anchor="center", width=40)
        self.tree_sales_history.column("qty", anchor="center", width=40)
        self.tree_sales_history.column("total", anchor="center", width=70)
        self.tree_sales_history.pack(fill="x", padx=10, pady=5)

        style = ttk.Style()
        style.theme_use("default")
        style.configure("Treeview", background=t["card"], foreground=t["text"], fieldbackground=t["card"], rowheight=28)
        style.configure("Treeview.Heading", background=t["bg"], foreground=t["accent"], font=("Segoe UI", 10, "bold"))
        
        self.update_goods_view(self.inventory)
        self.load_sales_history()

    def open_settings_window(self):
        t = self.themes[self.current_theme]
        l = self.lang_pack[self.current_lang]
        
        settings_win = tk.Toplevel(self.root)
        settings_win.title(l["settings_win"])
        settings_win.geometry("620x620")
        settings_win.configure(bg=t["bg"])
        settings_win.grab_set()
        
        sys_frame = tk.LabelFrame(settings_win, text=" ⚙️ System Customization / تخصيص النظام ", font=("Segoe UI", 10, "bold"), bg=t["card"], fg=t["accent"], bd=1)
        sys_frame.pack(fill="x", padx=20, pady=10, ipady=5)
        
        tk.Label(sys_frame, text=l["lang_label"], font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side="right" if self.current_lang=="ar" else "left", padx=5)
        combo_lang = ttk.Combobox(sys_frame, values=["العربية", "English"], width=10, state="readonly")
        combo_lang.set("العربية" if self.current_lang == "ar" else "English")
        combo_lang.pack(side="right" if self.current_lang=="ar" else "left", padx=5)
        
        tk.Label(sys_frame, text=l["theme_label"], font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side="right" if self.current_lang=="ar" else "left", padx=5)
        combo_thm = ttk.Combobox(sys_frame, values=["Dark", "Light", "Blue", "Green"], width=10, state="readonly")
        combo_thm.set(self.current_theme)
        combo_thm.pack(side="right" if self.current_lang=="ar" else "left", padx=5)
        
        def apply_changes():
            self.current_lang = "ar" if combo_lang.get() == "العربية" else "en"
            self.current_theme = combo_thm.get()
            self.save_config()
            settings_win.destroy()
            self.refresh_dashboard_ui()
            
        btn_apply = tk.Button(sys_frame, text="✅ Apply / تطبيق", bg="#2ECC71", fg="white", font=("Segoe UI", 9, "bold"), command=apply_changes)
        btn_apply.pack(fill="x", padx=10, pady=10)

        user_frame = tk.LabelFrame(settings_win, text=l["manage_users"], font=("Segoe UI", 10, "bold"), bg=t["card"], fg=t["accent"], bd=1)
        user_frame.pack(fill="x", padx=20, pady=10, ipady=5)
        
        f_user_row = tk.Frame(user_frame, bg=t["card"])
        f_user_row.pack(fill="x", padx=10, pady=5)
        
        side_align = "right" if self.current_lang == "ar" else "left"
        tk.Label(f_user_row, text=l["user_name"], font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        ent_username = tk.Entry(f_user_row, width=12)
        ent_username.pack(side=side_align, padx=5)
        
        tk.Label(f_user_row, text=l["new_pass"], font=("Segoe UI", 9), bg=t["card"], fg=t["text"]).pack(side=side_align, padx=2)
        ent_password = tk.Entry(f_user_row, width=12, show="*")
        ent_password.pack(side=side_align, padx=5)
        
        def add_or_update_user():
            u = ent_username.get().strip()
            p = ent_password.get().strip()
            if not u or not p: return
            self.users_db[u] = p
            self.save_users()
            messagebox.showinfo("Success", f"Saved: {u}", parent=settings_win)
            ent_username.delete(0, "end"); ent_password.delete(0, "end")
            
        def delete_user():
            u = ent_username.get().strip()
            if u in self.users_db and len(self.users_db) > 1:
                del self.users_db[u]
                self.save_users()
                messagebox.showinfo("Deleted", f"Removed: {u}", parent=settings_win)
                ent_username.delete(0, "end")

        btn_add_user = tk.Button(user_frame, text=l["save_user_btn"], bg="#2ECC71", fg="white", font=("Segoe UI", 9, "bold"), bd=0, command=add_or_update_user)
        btn_add_user.pack(fill="x", padx=10, pady=3)
        btn_del_user = tk.Button(user_frame, text=l["del_user_btn"], bg="#E74C3C", fg="white", font=("Segoe UI", 9, "bold"), bd=0, command=delete_user)
        btn_del_user.pack(fill="x", padx=10, pady=3)

        word_frame = tk.LabelFrame(settings_win, text=l["word_title"], font=("Segoe UI", 10, "bold"), bg=t["card"], fg="#FFD700", bd=1)
        word_frame.pack(fill="x", padx=20, pady=10, ipady=5)
        
        btn_import_word = tk.Button(word_frame, text=l["word_btn"], font=("Segoe UI", 11, "bold"), bg=t["btn_save"], fg="white", bd=0, height=2, command=lambda: self.import_from_word_action(settings_win))
        btn_import_word.pack(fill="x", padx=20, pady=15)

    def parse_clean_numeric(self, text, is_float=False):
        if not text:
            return 0.0 if is_float else 0
        cleaned = "".join(re.findall(r'[0-9.]' if is_float else r'[0-9]', str(text)))
        if not cleaned or cleaned == '.': 
            return 0.0 if is_float else 0
        try:
            return float(cleaned) if is_float else int(cleaned)
        except:
            return 0.0 if is_float else 0

    # 🚀 المحرك الخارق المعدل كلياً لتفكيك المدمج وتخطي عقبات بايثون مع النصوص العربية
    def import_from_word_action(self, window_context):
        if 'docx' not in globals(): 
            messagebox.showerror("Error", "python-docx package not installed!", parent=window_context)
            return
        file_path = filedialog.askopenfilename(title="Choose Word File", filetypes=[("Word Files", "*.docx")])
        if not file_path: return
        
        try:
            doc = docx.Document(file_path)
            count = 0
            print("\n--- بدأ فحص ملف الوورد السريع الحين ---")
            
            # 1. قراءة الجداول مع معالجة حرة للمستندات والدمج
            for idx, table in enumerate(doc.tables):
                print(f"جاري قراءة الجدول رقم {idx+1}...")
                for r_idx, row in enumerate(table.rows):
                    try:
                        # استخراج النصوص بدون حذف المتكرر الناتجة عن الدمج بشكل أعمى
                        raw_cells = [cell.text.strip() if cell.text else "" for cell in row.cells]
                        
                        # تصفية الخلايا الفارغة الحقيقية فقط
                        if not raw_cells or len(raw_cells) < 2:
                            continue
                            
                        # لو السطر يحتوي على كلمات عناوين زي كود أو صنف نتخطاه
                        if any(x in raw_cells[0] for x in ["كود", "code", "الكود"]):
                            continue
                        
                        # محاولة توزيع البيانات بمرونة
                        code = raw_cells[0] if raw_cells[0] else f"UNK-{count+1}"
                        name = raw_cells[1] if len(raw_cells) > 1 and raw_cells[1] else "صنف بدون اسم"
                        
                        price_raw = raw_cells[2] if len(raw_cells) > 2 else "0"
                        stock_raw = raw_cells[3] if len(raw_cells) > 3 else "0"
                        
                        price = self.parse_clean_numeric(price_raw, is_float=True)
                        stock = self.parse_clean_numeric(stock_raw, is_float=False)
                        
                        self.inventory[code] = {"name": name, "price": price, "stock": stock}
                        count += 1
                    except Exception as row_err:
                        print(f"ملحوظة بالسطر {r_idx} بالجدول: {row_err}")

            # 2. قراءة السطور العادية لو الجدول مش ظاهر كجدول حقيقي في بايثون
            if count == 0:
                print("لم يتم العثور على جداول متوافقة، جاري التبديل لقراءة الأسطر الحرة...")
                for p_idx, para in enumerate(doc.paragraphs):
                    line = para.text.strip()
                    if not line or any(x in line for x in ["كود", "code", "الجدول"]): 
                        continue
                    
                    # تقسيم السطر بجميع الفواصل الممكنة
                    parts = re.split(r'[,\t،\-_|]', line)
                    parts = [p.strip() for p in parts if p.strip()]
                    
                    if len(parts) >= 2:
                        code = parts[0]
                        name = parts[1]
                        price_raw = parts[2] if len(parts) > 2 else "0"
                        stock_raw = parts[3] if len(parts) > 3 else "0"
                        
                        price = self.parse_clean_numeric(price_raw, is_float=True)
                        stock = self.parse_clean_numeric(stock_raw, is_float=False)
                        
                        self.inventory[code] = {"name": name, "price": price, "stock": stock}
                        count += 1

            print(f"إجمالي الأصناف المقروءة بنجاح: {count}")
            self.save_goods(self.inventory)
            self.update_goods_view(self.inventory)
            
            if count > 0:
                messagebox.showinfo("Word Import", f"تم استيراد وحفظ {count} صنف بنجاح في قاعدة البيانات!", parent=window_context)
                window_context.destroy()
            else:
                messagebox.showwarning("تنبيه هوني", "تم قراءة الملف ولكن النتيجة 0 أصناف.\nتأكد أن النص في الوورد مرتب: كود ثم اسم ثم سعر ثم مخزن.", parent=window_context)
                
        except Exception as e:
            messagebox.showerror("Error", f"فشل قراءة الملف كلياً: {str(e)}", parent=window_context)

    def update_live_clock(self):
        current_time = time.strftime("⏱ %H:%M:%S | 📅 %Y-%m-%d")
        if hasattr(self, 'lbl_clock') and self.lbl_clock.winfo_exists():
            self.lbl_clock.config(text=current_time)
            self.root.after(1000, self.update_live_clock)

    def update_goods_view(self, data_dict):
        l = self.lang_pack[self.current_lang]
        for row in self.tree_goods.get_children(): self.tree_goods.delete(row)
        for code, info in data_dict.items():
            self.tree_goods.insert("", "end", values=(code, info["name"], f"{info['price']:.2f}{l['currency']}", info["stock"]))
        if hasattr(self, 'combo_size') and self.combo_size.winfo_exists():
            self.combo_size['values'] = list(data_dict.keys())
            if data_dict: self.combo_size.set(list(data_dict.keys())[0])

    def clean_text(self, text):
        text = text.lower().strip()
        replacements = {"ة": "ه", "أ": "ا", "إ": "ا", "آ": "ا", " ": ""}
        for old, new in replacements.items(): text = text.replace(old, new)
        return text

    def filter_inventory(self, event):
        query = self.clean_text(self.search_entry.get())
        filtered = {c: i for c, i in self.inventory.items() if query in self.clean_text(c) or query in self.clean_text(i["name"])}
        self.update_goods_view(filtered)

    def load_sales_history(self):
        for row in self.tree_sales_history.get_children(): self.tree_sales_history.delete(row)
        if os.path.exists(self.sales_file):
            with open(self.sales_file, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                try: next(reader)
                except StopIteration: return
                for row in reader:
                    if row and len(row) >= 7:
                        self.tree_sales_history.insert("", "end", values=(row[0], row[1], row[2], row[3], row[5], row[6]))

    def add_to_cart(self):
        size = self.combo_size.get()
        if not size: return
        qty_str = self.qty_entry.get()
        if not qty_str.isdigit() or int(qty_str) <= 0: return
        
        qty = int(qty_str)
        prod = self.inventory[size]
        already_in_cart = sum(item["qty"] for item in self.cart if item["code"] == size)
        
        if qty + already_in_cart > prod["stock"]: return
        total = qty * prod["price"]
        self.cart.append({"code": size, "name": prod["name"], "qty": qty, "price": prod["price"], "total": total})
        
        l = self.lang_pack[self.current_lang]
        for row in self.tree_cart.get_children(): self.tree_cart.delete(row)
        g_total = 0.0
        for item in self.cart:
            self.tree_cart.insert("", "end", values=(item["code"], item["name"], item["qty"], f"{item['total']:.2f}{l['currency']}"))
            g_total += item["total"]
        self.lbl_cart_total.config(text=l["invoice_total"]+f"{g_total:.2f}{l['currency']}")
        self.qty_entry.delete(0, "end")

    def checkout_cart(self):
        if not self.cart: return
        location = self.combo_location.get()
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        l = self.lang_pack[self.current_lang]
        
        for item in self.cart: self.inventory[item["code"]]["stock"] -= item["qty"]
        self.save_goods(self.inventory)
        
        with open(self.sales_file, "a", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            for item in self.cart:
                writer.writerow([now, location, self.current_user, item["code"], item["name"], item["qty"], f"{item['total']:.2f}{l['currency']}"])
                
        self.cart.clear()
        for row in self.tree_cart.get_children(): self.tree_cart.delete(row)
        self.lbl_cart_total.config(text=l["invoice_total"]+f"0.00{l['currency']}")
        self.update_goods_view(self.inventory)
        self.load_sales_history()

    def add_custom_product(self):
        code = self.new_code.get().strip()
        name = self.new_name.get().strip()
        price_str = self.new_price.get().strip()
        stock_str = self.new_stock.get().strip()
        if not (code and name and price_str and stock_str): return
        try:
            price = float(price_str)
            stock = int(stock_str)
        except ValueError: return
            
        self.inventory[code] = {"name": name, "price": price, "stock": stock}
        self.save_goods(self.inventory)
        self.update_goods_view(self.inventory)
        self.new_code.delete(0, "end"); self.new_name.delete(0, "end"); self.new_price.delete(0, "end"); self.new_stock.delete(0, "end")

if __name__ == "__main__":
    root = tk.Tk()
    app = CraftsmanProSystem(root)
    root.mainloop()
